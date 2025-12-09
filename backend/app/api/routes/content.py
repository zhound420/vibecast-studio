"""Content input API routes (upload, URL fetch)."""

import hashlib
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel, HttpUrl

from app.config import settings

router = APIRouter()


class URLFetchRequest(BaseModel):
    """Request schema for URL fetching."""
    url: HttpUrl
    extract_main_content: bool = True


class URLFetchResponse(BaseModel):
    """Response schema for URL fetching."""
    title: Optional[str]
    content: str
    word_count: int
    estimated_duration_minutes: int


class UploadResponse(BaseModel):
    """Response schema for file upload."""
    filename: str
    content: str
    word_count: int
    estimated_duration_minutes: int


@router.post("/url", response_model=URLFetchResponse)
async def fetch_from_url(request: URLFetchRequest):
    """
    Fetch content from a URL and extract readable text.
    Uses readability to extract main content.
    """
    import httpx
    from bs4 import BeautifulSoup

    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(
                str(request.url),
                follow_redirects=True,
                timeout=30.0,
                headers={
                    "User-Agent": "Mozilla/5.0 (compatible; VibeCast/1.0)"
                }
            )
            response.raise_for_status()
            html = response.text
    except httpx.HTTPError as e:
        raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {str(e)}")

    # Parse HTML and extract content
    soup = BeautifulSoup(html, "html.parser")

    # Get title
    title = None
    if soup.title:
        title = soup.title.string

    if request.extract_main_content:
        # Try to use readability for main content extraction
        try:
            from readability import Document
            doc = Document(html)
            content_html = doc.summary()
            content_soup = BeautifulSoup(content_html, "html.parser")
            content = content_soup.get_text(separator="\n", strip=True)
            if doc.title():
                title = doc.title()
        except Exception:
            # Fallback to basic extraction
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            content = soup.get_text(separator="\n", strip=True)
    else:
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        content = soup.get_text(separator="\n", strip=True)

    # Clean up content
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    content = "\n".join(lines)

    # Calculate stats
    word_count = len(content.split())
    estimated_duration = int(word_count / 150)  # ~150 WPM

    return URLFetchResponse(
        title=title,
        content=content,
        word_count=word_count,
        estimated_duration_minutes=estimated_duration,
    )


@router.post("/upload", response_model=UploadResponse)
async def upload_file(
    file: UploadFile = File(...),
):
    """
    Upload a text file and extract content.
    Supports: .txt, .md, .docx
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    # Check file extension
    suffix = Path(file.filename).suffix.lower()
    allowed_extensions = {".txt", ".md", ".docx", ".srt"}

    if suffix not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )

    # Read file content
    file_content = await file.read()

    try:
        if suffix in {".txt", ".md", ".srt"}:
            # Plain text files
            content = file_content.decode("utf-8")
        elif suffix == ".docx":
            # Word documents
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            content = "\n".join([para.text for para in doc.paragraphs])
        else:
            content = file_content.decode("utf-8")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")

    # Calculate stats
    word_count = len(content.split())
    estimated_duration = int(word_count / 150)  # ~150 WPM

    return UploadResponse(
        filename=file.filename,
        content=content,
        word_count=word_count,
        estimated_duration_minutes=estimated_duration,
    )


class SegmentationRequest(BaseModel):
    """Request schema for content segmentation."""
    content: str
    num_speakers: int = 2
    style: str = "auto"  # auto, paragraph, sentence


class SegmentationResponse(BaseModel):
    """Response schema for segmentation."""
    segments: list[dict]
    speaker_count: int
    total_segments: int


@router.post("/segment", response_model=SegmentationResponse)
async def segment_content(request: SegmentationRequest):
    """
    Intelligently segment content and assign speakers.
    """
    from app.services.content.segmenter import ContentSegmenter

    segmenter = ContentSegmenter()
    segments = segmenter.segment(
        content=request.content,
        num_speakers=request.num_speakers,
        style=request.style,
    )

    # Count unique speakers
    speaker_ids = set(seg["speaker_id"] for seg in segments)

    return SegmentationResponse(
        segments=segments,
        speaker_count=len(speaker_ids),
        total_segments=len(segments),
    )
